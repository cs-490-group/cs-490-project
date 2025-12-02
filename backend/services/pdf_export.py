"""
Export salary negotiation materials to PDF and DOCX formats
UC-083: Export research summary for offline preparation
"""

from io import BytesIO
from datetime import datetime
from typing import Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

async def export_negotiation_to_pdf(offer: Dict[str, Any]) -> BytesIO:
    """Export negotiation preparation to PDF format"""
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=f"Salary Negotiation Prep - {offer['company']}"
    )

    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f4788'),
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2e5c8a'),
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6
    )

    # Title
    story.append(Paragraph(f"Salary Negotiation Preparation", title_style))
    story.append(Spacer(1, 0.1 * inch))

    prep = offer.get('negotiation_prep', {})

    # Header Info
    header_data = [
        ['Position:', offer.get('job_title', 'N/A')],
        ['Company:', offer.get('company', 'N/A')],
        ['Location:', offer.get('location', 'N/A')],
        ['Offered Salary:', f"${offer.get('offered_salary_details', {}).get('base_salary', 0):,}"],
    ]

    header_table = Table(header_data, colWidths=[1.5 * inch, 4 * inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8f0f7')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
    ]))

    story.append(header_table)
    story.append(Spacer(1, 0.2 * inch))

    # Executive Summary
    if prep.get('executive_summary'):
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = prep['executive_summary'].replace('\n', '<br/>')
        story.append(Paragraph(summary_text, normal_style))
        story.append(Spacer(1, 0.1 * inch))

    # Market Salary Data
    if prep.get('market_salary_data'):
        story.append(Paragraph("Market Salary Research", heading_style))
        market = prep['market_salary_data']

        market_data = [
            ['Metric', 'Value'],
            ['Median Salary', f"${market.get('median_salary', 0):,}"],
            ['25th Percentile', f"${market.get('percentile_25', 0):,}"],
            ['75th Percentile', f"${market.get('percentile_75', 0):,}"],
            ['90th Percentile', f"${market.get('percentile_90', 0):,}"],
            ['Industry Average', f"${market.get('industry_average', 0):,}"],
            ['Salary Trend', market.get('salary_trend', 'N/A')],
        ]

        market_table = Table(market_data, colWidths=[2 * inch, 2.5 * inch])
        market_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e5c8a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
        ]))

        story.append(market_table)
        story.append(Spacer(1, 0.2 * inch))

    # Best Practices
    if prep.get('best_practices'):
        story.append(PageBreak())
        story.append(Paragraph("Best Practices", heading_style))

        for i, practice in enumerate(prep['best_practices'][:10], 1):
            story.append(Paragraph(f"<b>{i}.</b> {practice}", normal_style))

        story.append(Spacer(1, 0.1 * inch))

    # Generated timestamp
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        f"<i>Generated on {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')}</i>",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
    ))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


async def export_negotiation_to_docx(offer: Dict[str, Any]) -> BytesIO:
    """Export negotiation preparation to DOCX format"""
    doc = Document()

    # Title
    title = doc.add_heading(f"Salary Negotiation Preparation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header Info
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Position:'
    header_cells[1].text = offer.get('job_title', 'N/A')

    table.rows[1].cells[0].text = 'Company:'
    table.rows[1].cells[1].text = offer.get('company', 'N/A')

    table.rows[2].cells[0].text = 'Location:'
    table.rows[2].cells[1].text = offer.get('location', 'N/A')

    table.rows[3].cells[0].text = 'Offered Salary:'
    table.rows[3].cells[1].text = f"${offer.get('offered_salary_details', {}).get('base_salary', 0):,}"

    doc.add_paragraph()

    prep = offer.get('negotiation_prep', {})

    # Executive Summary
    if prep.get('executive_summary'):
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(prep['executive_summary'])

    # Market Salary Data
    if prep.get('market_salary_data'):
        doc.add_heading('Market Salary Research', level=1)
        market = prep['market_salary_data']

        market_table = doc.add_table(rows=8, cols=2)
        market_table.style = 'Light Grid Accent 1'

        market_table.rows[0].cells[0].text = 'Metric'
        market_table.rows[0].cells[1].text = 'Value'

        market_table.rows[1].cells[0].text = 'Median Salary'
        market_table.rows[1].cells[1].text = f"${market.get('median_salary', 0):,}"

        market_table.rows[2].cells[0].text = '25th Percentile'
        market_table.rows[2].cells[1].text = f"${market.get('percentile_25', 0):,}"

        market_table.rows[3].cells[0].text = '75th Percentile'
        market_table.rows[3].cells[1].text = f"${market.get('percentile_75', 0):,}"

        market_table.rows[4].cells[0].text = '90th Percentile'
        market_table.rows[4].cells[1].text = f"${market.get('percentile_90', 0):,}"

        market_table.rows[5].cells[0].text = 'Industry Average'
        market_table.rows[5].cells[1].text = f"${market.get('industry_average', 0):,}"

        market_table.rows[6].cells[0].text = 'Salary Trend'
        market_table.rows[6].cells[1].text = market.get('salary_trend', 'N/A')

        market_table.rows[7].cells[0].text = 'Company Size'
        market_table.rows[7].cells[1].text = market.get('company_size_factor', 'N/A')

        doc.add_paragraph()

    # Best Practices
    if prep.get('best_practices'):
        doc.add_heading('Best Practices', level=1)

        for i, practice in enumerate(prep['best_practices'][:15], 1):
            doc.add_paragraph(f"{practice}", style='List Number')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Generated on {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
