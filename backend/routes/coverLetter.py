from fastapi import APIRouter, HTTPException, Header, Path, File, UploadFile, Form, Body, Depends
from fastapi.responses import FileResponse
from typing import List, Optional, Dict
from datetime import datetime
from uuid import uuid4
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import tempfile
import requests
from xhtml2pdf import pisa
import re
from bs4 import BeautifulSoup, NavigableString, Tag

from schema.CoverLetter import CoverLetterIn, CoverLetterOut, CoverLetterUpdate,CoverLetterShare, CoverLetterFeedback,ApprovalRequest,CoverLetterVersion
from mongo.cover_letters_dao import cover_letters_dao
from mongo.jobs_dao import jobs_dao
from sessions.session_authorizer import authorize

coverletter_router = APIRouter(prefix="/cover-letters")


class NativeDocxGenerator:
    def __init__(self, html_content):
        self.doc = Document()
        # Set standardized margins
        section = self.doc.sections[0]
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        
        self.soup = BeautifulSoup(html_content, 'html.parser')

    def convert(self):
        # 1. Background Color extraction (Best effort)
        bg_color = self._extract_bg_color(self.soup.body)
        if bg_color:
            self._set_page_background(bg_color)

        # 2. Find Container
        container = self.soup.find("div", class_="container")
        
        if container:
            self._process_container_as_table(container)
        else:
            # Fallback for no container
            root = self.soup.body if self.soup.body else self.soup
            self._process_node(root, self.doc, {})
            
        return self.doc

    def _process_container_as_table(self, container):
        """Renders the main container as a Centered Table with Borders."""
        # Create a table (1x1)
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        
        # Set Width to 90% of page (approx 4500 pct units)
        tbl_pr = table._element.tblPr
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), '4500') 
        tbl_w.set(qn('w:type'), 'pct')
        tbl_pr.append(tbl_w)

        # ADD BORDERS (To mimic the 'card' look)
        self._set_table_borders(table)

        cell = table.cell(0, 0)
        
        # Add Padding inside the cell
        # (This mimics CSS padding)
        tc_pr = cell._element.get_or_add_tcPr()
        tc_mar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom', 'left', 'right']:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), '280') # ~0.5cm padding
            node.set(qn('w:type'), 'dxa')
            tc_mar.append(node)
        tc_pr.append(tc_mar)

        # Background Color
        styles = self._parse_styles(container.get('style', ''))
        bg_hex = "FFFFFF" # Default white card
        if "background" in styles and "#" in styles["background"]:
             match = re.search(r'#(?:[0-9a-fA-F]{3}){1,2}', styles["background"])
             if match: bg_hex = match.group(0).replace("#", "")
        
        self._set_cell_shading(cell, bg_hex)
        
        cell._element.clear_content()
        
        # Process children
        for child in container.children:
            self._process_node(child, cell, styles)

    def _process_node(self, element, parent, inherited_styles: Dict, current_paragraph=None):
        if isinstance(element, NavigableString):
            text = str(element)
            if not text.strip() and not current_paragraph: return
            if current_paragraph is None: current_paragraph = parent.add_paragraph()
            
            run = current_paragraph.add_run(text)
            self._apply_run_formatting(run, inherited_styles)
            return

        if not isinstance(element, Tag): return

        tag = element.name.lower()
        styles = self._parse_styles(element.get('style', ''))
        
        current_styles = inherited_styles.copy()
        current_styles.update(styles)
        if tag in ['b', 'strong']: current_styles['font-weight'] = 'bold'
        if tag in ['i', 'em']: current_styles['font-style'] = 'italic'
        if tag == 'u': current_styles['text-decoration'] = 'underline'

        # Blocks
        if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'div', 'li']:
            p = parent.add_paragraph()
            
            # Alignment
            align = current_styles.get('text-align')
            if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif 'header' in element.get('class', []): p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
            
            # Headers
            if tag == 'h1': p.style = 'Heading 1'
            if tag == 'h2': p.style = 'Heading 2'
            
            # Spacing
            p.paragraph_format.space_after = Pt(12) 
            current_paragraph = p
            
        elif tag == 'br':
            if current_paragraph: current_paragraph.add_run().add_break()
            return

        for child in element.children:
            self._process_node(child, parent, current_styles, current_paragraph)

    # --- XML Helpers ---

    def _set_table_borders(self, table):
        """Adds a thin border to the table to simulate a container box."""
        tbl_pr = table._element.tblPr
        tbl_borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') # 4 = 1/2 pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'DDDDDD') # Light gray border
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def _set_cell_shading(self, cell, hex_color):
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    def _set_page_background(self, hex_color):
        background = OxmlElement('w:background')
        background.set(qn('w:color'), hex_color)
        self.doc.element.insert(0, background)
        display = OxmlElement('w:displayBackgroundShape')
        self.doc.settings.element.append(display)

    def _parse_styles(self, style_str: str) -> Dict:
        if not style_str: return {}
        return {k.strip().lower(): v.strip() for k, v in [x.split(':', 1) for x in style_str.split(';') if ':' in x]}

    def _extract_bg_color(self, element):
        if not element: return None
        styles = self._parse_styles(element.get('style', ''))
        bg = styles.get('background', '') or styles.get('background-color', '')
        match = re.search(r'#(?:[0-9a-fA-F]{3}){1,2}', bg)
        return match.group(0).replace("#", "") if match else None

    def _apply_run_formatting(self, run, styles):
        if 'bold' in styles.get('font-weight', ''): run.font.bold = True
        if 'italic' in styles.get('font-style', ''): run.font.italic = True
        if 'underline' in styles.get('text-decoration', ''): run.font.underline = True
        
        color = styles.get('color')
        if color and '#' in color:
            hex_code = re.search(r'#(?:[0-9a-fA-F]{3}){1,2}', color)
            if hex_code:
                h = hex_code.group(0).lstrip('#')
                if len(h) == 3: h = ''.join([c*2 for c in h])
                run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))



# ============================================================
# GET usage stats aggregated by template type
# ============================================================
@coverletter_router.get("/usage/by-type")
async def get_usage_by_template_type():
    """Get aggregated usage counts grouped by template type (style_industry)."""
    try:
        usage_stats = await cover_letters_dao.get_usage_by_template_type()
        return usage_stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get usage stats: {str(e)}")

# ============================================================
# GET all cover letters for the current user
# ============================================================
@coverletter_router.get("/me", response_model=List[CoverLetterOut])
async def get_my_coverletters(uuid: str = Depends(authorize)):
    """Fetch all cover letters belonging to the current user."""
    letters = await cover_letters_dao.get_all_cover_letters(uuid)

    if not letters:
        return []

    mapped_letters = [
        {
            "id": str(l["_id"]),
            "user_id": l.get("uuid"),
            "title": l.get("title"),
            "company": l.get("company"),
            "position": l.get("position"),
            "content": l.get("content"),
            "created_at": l.get("created_at"),
            "usage_count": l.get("usage_count", 0),
            "default_cover_letter": l.get("default_cover_letter", False),
            "job_id":l.get("job_id")
        }
        for l in letters
    ]

    return mapped_letters

# ============================================================
# GET a single cover letter by ID
# ============================================================
@coverletter_router.get("/{letter_id}", response_model=CoverLetterOut)
async def get_coverletter(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Fetch a single cover letter by ID if it belongs to the current user."""
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)

    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")

    return {
        "id": str(letter["_id"]),
        "user_id": letter.get("uuid"),
        "title": letter.get("title"),
        "company": letter.get("company"),
        "position": letter.get("position"),
        "content": letter.get("content"),
        "created_at": letter.get("created_at"),
        "usage_count": letter.get("usage_count", 0),
        "default_cover_letter": letter.get("default_cover_letter", False),
        "job_id": letter.get("job_id")
    }

# ============================================================
# POST create a new cover letter
# ============================================================
@coverletter_router.post("")
async def add_coverletter(
    coverletter: CoverLetterIn,
    uuid: str = Depends(authorize)
):
    """Add a new cover letter for the current user."""
    new_letter = {
        "_id": str(uuid4()),
        "uuid": uuid,
        "title": coverletter.title,
        "company": coverletter.company,
        "position": coverletter.position,
        "content": coverletter.content,
        "created_at": datetime.utcnow().isoformat(),
        "usage_count": 1 if coverletter.template_type else 0,
        "template_type": getattr(coverletter, 'template_type', None),
        "default_cover_letter": False,
        "job_id": ""
    }

    inserted_id = await cover_letters_dao.add_cover_letter(new_letter)
    
    return {
        "coverletter_id": inserted_id,
        "data": {
            "_id": inserted_id,
            "title": new_letter["title"],
            "created_at": new_letter["created_at"]
        }
    }

# ============================================================
# POST upload an HTML/file cover letter
# ============================================================
@coverletter_router.post("/upload")
async def upload_coverletter(
    file: UploadFile = File(...),
    title: str = Form(...),
    company: str = Form(""),
    position: str = Form(""),
    version_name: str = Form("Version 1"),
    description: str = Form(None),
    uuid: str = Depends(authorize)
):
    """Upload a cover letter file (HTML, PDF, or DOCX)."""
    
    if file.content_type != "text/html" and not file.filename.lower().endswith('.html'):
        raise HTTPException(status_code=400, detail="Only HTML files are supported currently")
    
    try:
        content = await file.read()
        html_content = content.decode('utf-8')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")
    
    letter_id = str(uuid4())
    created_at = datetime.utcnow().isoformat()
    
    new_letter = {
        "_id": letter_id,
        "uuid": uuid,
        "title": title,
        "company": company,
        "position": position,
        "content": html_content,
        "version_name": version_name,
        "description": description,
        "created_at": created_at,
        "usage_count": 0,
        "uploadedFile": True,
        "template_type": None,
        "file_type": file.content_type,
        "file_size": len(content),
        "default_cover_letter": False,
        "job_id" : ""
    }
    
    await cover_letters_dao.add_cover_letter(new_letter)
    
    return {
        "detail": "File uploaded successfully",
        "data": {
            "_id": letter_id,
            "title": title,
            "version_name": version_name,
            "created_at": created_at,
            "file_size": len(content)
        }
    }

# ============================================================
# PUT update a cover letter
# ============================================================
@coverletter_router.put("/{letter_id}")
async def update_coverletter(
    letter_id: str = Path(...),
    coverletter: CoverLetterUpdate = Body(...), 
    uuid: str = Depends(authorize)
):
    updates = coverletter.model_dump(exclude_unset=True)
    
    if not updates:
        return {"message": "No fields to update"}

    modified_count = await cover_letters_dao.update_cover_letter(letter_id, uuid, updates)

    if modified_count == 0:
        letter_exists = await cover_letters_dao.get_cover_letter(letter_id, uuid)
        if not letter_exists:
            raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")
        return {"message": "No changes to update"}

    return {"message": "Updated successfully"}

# ============================================================
# DELETE a cover letter
# ============================================================
@coverletter_router.delete("/{letter_id}")
async def delete_coverletter(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")
    
    deleted_count = await cover_letters_dao.delete_cover_letter(letter_id)

    if deleted_count == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")

    return {"message": "Deleted successfully"}

# ============================================================
# PUT set as default cover letter
# ============================================================
@coverletter_router.put("/{letter_id}/default")
async def set_default_coverletter(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found or not owned by user")
    
    updated = await cover_letters_dao.set_default_cover_letter(letter_id, uuid)
    
    if updated == 0:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    return {"message": "Default cover letter set successfully"}

# ============================================================
# GET download as PDF
# ============================================================
@coverletter_router.get("/{letter_id}/download/pdf")
async def download_pdf(
    letter_id: str = Path(...),
    uuid: str = Depends(authorize)
):
    """Download cover letter as PDF - self-hosted generation"""
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    
    try:
        html_content = letter.get("content", "")
        filename = f"{letter.get('title', 'cover_letter')}.pdf"
        
        # Generate PDF using xhtml2pdf
        pdf_buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            io.BytesIO(html_content.encode('utf-8')),
            dest=pdf_buffer
        )
        
        if pisa_status.err:
            raise HTTPException(status_code=500, detail="PDF generation failed")
        
        pdf_buffer.seek(0)
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(pdf_buffer.getvalue())
            tmp_path = tmp.name
        
        return FileResponse(
            tmp_path,
            media_type='application/pdf',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    except Exception as e:
        print(f"PDF Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

# ============================================================
# GET download as DOCX (UPGRADED for Styling)
# ============================================================
@coverletter_router.get("/{letter_id}/download/docx")
async def download_docx(letter_id: str = Path(...), uuid: str = Depends(authorize)):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter: raise HTTPException(404)
    
    try:
        html_content = letter.get("content", "").replace("&nbsp;", " ")
        generator = NativeDocxGenerator(html_content)
        doc = generator.convert()
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        filename = f"{letter.get('title', 'letter').replace(' ', '_')}.docx"
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(doc_buffer.getvalue())
            tmp_path = tmp.name
            
        return FileResponse(tmp_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename, headers={"Content-Disposition": f"attachment; filename=\"{filename}\""})
    except Exception as e:
        print(f"DOCX Gen Error: {e}")
        raise HTTPException(500, f"Error: {e}")

# ============================================================
# HELPER FUNCTIONS: Advanced HTML to DOCX Processing
# ============================================================

def parse_style(style_str):
    """Parse inline CSS style string into a dictionary."""
    if not style_str:
        return {}
    styles = {}
    for item in style_str.split(';'):
        if ':' in item:
            key, value = item.split(':', 1)
            styles[key.strip().lower()] = value.strip()
    return styles

def hex_to_rgb(hex_color):
    """Convert hex color string to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    try:
        if len(hex_color) == 3:
            hex_color = ''.join([c*2 for c in hex_color])
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except:
        return None

def apply_run_formatting(run, styles, tag_name):
    """Apply formatting to a docx run based on CSS styles and tag."""
    # Bold / Italic / Underline
    if tag_name in ['b', 'strong'] or 'bold' in styles.get('font-weight', ''):
        run.font.bold = True
    if tag_name in ['i', 'em'] or 'italic' in styles.get('font-style', ''):
        run.font.italic = True
    if tag_name == 'u' or 'underline' in styles.get('text-decoration', ''):
        run.font.underline = True
        
    # Color
    color = styles.get('color')
    if color and color.startswith('#'):
        rgb = hex_to_rgb(color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
            
    # Font Size
    size_str = styles.get('font-size')
    if size_str:
        try:
            val = float(re.sub(r'[^\d.]', '', size_str))
            if 'pt' in size_str:
                run.font.size = Pt(val)
            elif 'px' in size_str:
                # Approx conversion: 1px = 0.75pt
                run.font.size = Pt(val * 0.75) 
        except:
            pass

def process_html_node(element, doc, current_paragraph=None):
    """
    Recursively process HTML elements.
    - Block elements (p, div, h1-h6) create new paragraphs.
    - Inline elements (span, b, i) append runs to current_paragraph.
    """
    if isinstance(element, NavigableString):
        text = str(element)
        # Avoid adding runs for pure whitespace between block elements
        if current_paragraph is None and not text.strip():
            return
            
        # If text exists but no paragraph context, create one (fallback)
        if current_paragraph is None:
            current_paragraph = doc.add_paragraph()
            
        current_paragraph.add_run(text)
        return

    if element.name is None:
        return

    tag = element.name.lower()
    styles = parse_style(element.get('style', ''))
    
    # --- BLOCK ELEMENTS ---
    if tag in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'header', 'footer', 'section', 'article']:
        p = doc.add_paragraph()
        
        # Headers
        if tag == 'h1': p.style = 'Heading 1'
        elif tag == 'h2': p.style = 'Heading 2'
        elif tag == 'h3': p.style = 'Heading 3'
        
        # Alignment
        align = styles.get('text-align')
        if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # List Items
        if tag == 'li':
            p.style = 'List Bullet'
            
        # Recursive processing for children with NEW paragraph context
        for child in element.children:
            process_html_node(child, doc, p)
            
    # --- INLINE ELEMENTS ---
    elif tag in ['span', 'b', 'strong', 'i', 'em', 'u', 'a', 'small', 'mark']:
        if current_paragraph is None:
            current_paragraph = doc.add_paragraph()
            
        # For inline elements, we need to apply formatting to specific text runs.
        # Since we can't easily nest runs, we iterate children.
        # Note: This simple recursion applies formatting to immediate text children.
        for child in element.children:
            if isinstance(child, NavigableString):
                run = current_paragraph.add_run(str(child))
                apply_run_formatting(run, styles, tag)
            else:
                # If nested inline tag (e.g. <b><i>text</i></b>), recurse
                # Ideally we merge styles here, but simple recursion works for common cases
                process_html_node(child, doc, current_paragraph)

    elif tag == 'br':
        if current_paragraph:
            current_paragraph.add_run().add_break()
            
    elif tag == 'ul' or tag == 'ol':
        # Lists are containers; process children (LIs will make their own paragraphs)
        for child in element.children:
            process_html_node(child, doc, None)
            
    else:
        # Fallback container (process children)
        for child in element.children:
            process_html_node(child, doc, current_paragraph)


# ============================================================
# (KEEP EXISTING PUBLIC/SHARE/FEEDBACK/VERSION ENDPOINTS)
# ============================================================

@coverletter_router.get("/public/{token}", tags=["cover-letters"])
async def get_shared_cover_letter(token: str):
    try:
        result = await cover_letters_dao.get_cover_letter_by_token(token)
        if result:
            return result
        raise HTTPException(400, "Invalid or expired share link")
    except Exception as e:
        raise HTTPException(500, f"Error: {str(e)}")

@coverletter_router.post("/public/{token}/feedback", tags=["cover-letters"])
async def add_public_feedback(token: str, feedback: CoverLetterFeedback):
    try:
        letter = await cover_letters_dao.get_cover_letter_by_token(token)
        if not letter:
            raise HTTPException(400, "Invalid token")
            
        settings = letter.get("share_settings", {})
        if not settings.get("can_comment"):
            raise HTTPException(403, "Comments are disabled for this document")

        model = feedback.model_dump()
        model["cover_letter_id"] = letter["_id"]
        
        feedback_id = await cover_letters_dao.add_feedback(model)
        return {"detail": "Feedback added", "feedback_id": feedback_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))

@coverletter_router.post("/{letter_id}/share", tags=["cover-letters"])
async def create_share_link(letter_id: str, share: CoverLetterShare, uuid: str = Depends(authorize)):
    try:
        model = share.model_dump()
        result = await cover_letters_dao.create_share_link(letter_id, model)
        return {"detail": "Link generated", "share_link": result["token"], "share_data": result}
    except Exception as e:
        raise HTTPException(500, str(e))

@coverletter_router.get("/{letter_id}/share", tags=["cover-letters"])
async def get_share_link(letter_id: str, uuid: str = Depends(authorize)):
    try:
        result = await cover_letters_dao.get_share_link(letter_id)
        if not result:
            raise HTTPException(400, "No active link found")
        return result
    except Exception as e:
        raise HTTPException(500, str(e))

@coverletter_router.delete("/{letter_id}/share", tags=["cover-letters"])
async def revoke_share_link(letter_id: str, uuid: str = Depends(authorize)):
    await cover_letters_dao.revoke_share_link(letter_id)
    return {"detail": "Link revoked"}

@coverletter_router.get("/{letter_id}/feedback", tags=["cover-letters"])
async def get_feedback(letter_id: str, uuid: str = Depends(authorize)):
    return await cover_letters_dao.get_feedback(letter_id)

@coverletter_router.post("/{letter_id}/feedback", tags=["cover-letters"])
async def add_internal_feedback(letter_id: str, feedback: CoverLetterFeedback, uuid: str = Depends(authorize)):
    model = feedback.model_dump()
    model["cover_letter_id"] = letter_id
    res = await cover_letters_dao.add_feedback(model)
    return {"detail": "Feedback added", "feedback_id": res}

@coverletter_router.put("/{letter_id}/feedback/{feedback_id}", tags=["cover-letters"])
async def update_feedback(letter_id: str, feedback_id: str, feedback: CoverLetterFeedback, uuid: str = Depends(authorize)):
    model = feedback.model_dump(exclude_unset=True)
    await cover_letters_dao.update_feedback(feedback_id, model)
    return {"detail": "Feedback updated"}

@coverletter_router.delete("/{letter_id}/feedback/{feedback_id}", tags=["cover-letters"])
async def delete_feedback(letter_id: str, feedback_id: str, uuid: str = Depends(authorize)):
    await cover_letters_dao.delete_feedback(feedback_id)
    return {"detail": "Feedback deleted"}

@coverletter_router.post("/public/{token}/status", tags=["cover-letters"])
async def set_public_cover_letter_status(token: str, request: ApprovalRequest):
    letter = await cover_letters_dao.get_cover_letter_by_token(token)
    if not letter:
        raise HTTPException(400, "Invalid token")
        
    if not letter.get("share_settings", {}).get("can_comment"):
        raise HTTPException(403, "Reviewer permissions disabled")

    await cover_letters_dao.update_approval_status(letter["_id"], request.status)
    return {"detail": f"Status updated to {request.status}"}

@coverletter_router.post("/{letter_id}/versions", tags=["cover-letters"])
async def create_version(letter_id: str, version: CoverLetterVersion, uuid: str = Depends(authorize)):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(404, "Cover letter not found")

    model = version.model_dump()
    if not model.get("content_snapshot"):
        model["content_snapshot"] = letter.get("content")
    if not model.get("title_snapshot"):
        model["title_snapshot"] = letter.get("title")
        
    result = await cover_letters_dao.create_version(letter_id, model)
    return {"detail": "Version saved", "version_id": result}

@coverletter_router.get("/{letter_id}/versions", tags=["cover-letters"])
async def get_versions(letter_id: str, uuid: str = Depends(authorize)):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(404, "Cover letter not found")
        
    return await cover_letters_dao.get_versions(letter_id)

@coverletter_router.post("/{letter_id}/versions/{version_id}/restore", tags=["cover-letters"])
async def restore_version(letter_id: str, version_id: str, uuid: str = Depends(authorize)):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(404, "Cover letter not found")

    updated = await cover_letters_dao.restore_version(letter_id, version_id)
    if updated == 0:
        raise HTTPException(400, "Failed to restore version")
    return {"detail": "Version restored successfully"}

@coverletter_router.delete("/{letter_id}/versions/{version_id}", tags=["cover-letters"])
async def delete_version(letter_id: str, version_id: str, uuid: str = Depends(authorize)):
    letter = await cover_letters_dao.get_cover_letter(letter_id, uuid)
    if not letter:
        raise HTTPException(404, "Cover letter not found")

    await cover_letters_dao.delete_version(version_id)
    return {"detail": "Version deleted"}

@coverletter_router.get("/analytics/performance", tags=["cover-letters"])
async def get_cover_letter_analytics(uuid: str = Depends(authorize)):
    try:
        stats = await cover_letters_dao.get_performance_stats(uuid, jobs_dao)
        insights = []
        styles = stats.get("styles", [])
        
        if styles:
            best_style = max(styles, key=lambda x: x["interview_rate"])
            if best_style["interview_rate"] > 0:
                insights.append(f"🏆 Your '{best_style['style']}' letters have the highest interview rate ({best_style['interview_rate']}%)")
            
            if len(styles) > 1:
                worst_style = min(styles, key=lambda x: x["interview_rate"])
                diff = best_style["interview_rate"] - worst_style["interview_rate"]
                if diff > 10:
                    insights.append(f"💡 Switching from '{worst_style['style']}' to '{best_style['style']}' could boost your interviews by {diff:.1f}%")
        
        stats["insights"] = insights
        return stats
        
    except Exception as e:
        print(f"Analytics Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate analytics")